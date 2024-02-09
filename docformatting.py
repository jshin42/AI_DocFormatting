import os
import docx
import openai

# Define your OpenAI API key
API_KEY = 'YOUR_OPENAPI_KEY'

# Function to read a Word document and return its content
def read_docx(filename):
    doc = docx.Document(filename)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def enhance_text(text, api_key):
    openai.api_key = api_key
    prompt = (
        "Please clean up the following doc, remove extra urls and images. "
        "Make the following text more readable and concise. "
        "No corporate jargon, explain to me like I'm a software architect, "
        "make the audience tailored towards a product manager.\n"
        f"{text}\n"
    )
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Use the text-based model
        messages=[
            {"role": "system", "content": "Prompt:"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500  # Adjust as needed
    )
    # Assuming the last message in the conversation is the AI's completion
    if response and 'choices' in response and len(response['choices']) > 0:
        last_message = response['choices'][0]['message']['content']
        return last_message.strip()
    else:
        # Handle error or unexpected response format
        return "Could not process the text, please try again."



# Function to write enhanced text to a Word document
def write_docx(filename, enhanced_text):
    doc = docx.Document()
    doc.add_paragraph(enhanced_text)
    doc.save(filename)

# Function to split input document into smaller chunks
def split_document(input_filename, output_folder):
    doc = docx.Document(input_filename)
    content = [p.text for p in doc.paragraphs]
    chunk_size = 500  # Adjust chunk size as needed
    for i in range(0, len(content), chunk_size):
        chunk = "\n".join(content[i:i+chunk_size])
        output_filename = os.path.join(output_folder, f"output_{i//chunk_size}.docx")
        write_docx(output_filename, chunk)

# Function to combine output documents into one large document
def combine_documents(input_folder, output_filename):
    output_doc = docx.Document()
    for filename in sorted(os.listdir(input_folder)):
        if filename.endswith(".docx"):
            filepath = os.path.join(input_folder, filename)
            content = read_docx(filepath)
            output_doc.add_paragraph(content)
    output_doc.save(output_filename)

# Input and output filenames
INPUT_FILENAME = "input.docx"
OUTPUT_FOLDER = "output"
COMBINED_OUTPUT_FILENAME = "combined_output.docx"

# Split input document into smaller chunks
split_document(INPUT_FILENAME, OUTPUT_FOLDER)

# Enhance each chunk separately
for filename in os.listdir(OUTPUT_FOLDER):
    if filename.endswith(".docx"):
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        content = read_docx(filepath)
        enhanced_content = enhance_text(content, API_KEY)
        write_docx(filepath, enhanced_content)

# Combine enhanced documents into one large document
combine_documents(OUTPUT_FOLDER, COMBINED_OUTPUT_FILENAME)

print("Documents processed and combined.")
